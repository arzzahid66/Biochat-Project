import io
from typing import List, TypedDict
from dotenv import load_dotenv
load_dotenv()
from langchain_openai import ChatOpenAI
from openai import OpenAI
from app.utils.utils_code import encode_image, namespace_generator, IncomingFileProcessor
from app.utils.lcel_chain_pinecone_class import SimpleQAChains, PineconeInsertRetrieval
from langchain_openai import OpenAIEmbeddings
from langgraph.graph import StateGraph, START, END
from langchain_core.documents import Document as LangchainDocument
import os
import re
import fitz  # PyMuPDF
from io import BytesIO
from docx import Document
# Ensure you have python-docx installed

client = OpenAI()
Model = "gpt-4o-mini"

model = ChatOpenAI(model="gpt-4o-mini")
collection_name = "BioChat"
# Embeddings and classes 
embeddings = OpenAIEmbeddings(model="text-embedding-3-small")
file_processor = IncomingFileProcessor()
pinecone_key = os.getenv("PINECONE_API_KEY")
pinecone_index_name = os.getenv("PINECONE_INDEX_NAME")
pinecone_ar_class = PineconeInsertRetrieval(api_key=pinecone_key)
archains = SimpleQAChains(model=model)

client = OpenAI()
llm = ChatOpenAI(model_name="gpt-4o-mini")

class State(TypedDict):
    output_result: str
    image_bytes: bytes
    file_text : str
    original_file_name: str
    excel_csv_file_key: bytes

# pandas df bytes key 
def pandas_df_agent(state):
    file_key = state["excel_csv_file_key"]
    return {"output_result": file_key}

# QA retrieval
def QA_Retriever_node(state):
    try:
        original_file_name = state["original_file_name"]
        file_raw_text = state["file_text"]
        file_extension = os.path.splitext(original_file_name)[1].lower()
        # Get documents based on file type
        sorted_extention = [".docx", ".pdf", ".doc", ".txt", ".png", ".jpg", ".jpeg"]
        if file_extension in sorted_extention:
            docs = file_processor.preprocess_files_all(file_raw_text, original_file_name)
        else:   
            raise ValueError("Unsupported file extension")
        if not docs:
            return {"output_result": "No content could be extracted from the document"}
        
        # Generate namespace with 4 random words and current datetime
        name_space = namespace_generator()
        try:
            pinecone_ar_class.insert_data_in_namespace(docs, embeddings, pinecone_index_name, name_space)
        except Exception as e:
            return {"output_result": f"An error occurred while processing the document: {str(e)}"}
        return {"output_result": name_space}
    except Exception as e:
        print(f"Error in QA_Retriever_node: {str(e)}")
        return {"output_result": f"An error occurred while processing the document: {str(e)}"}

# image qa
def image_qa(state):
    try:    
        image_bytes = state["image_bytes"]
        base64_image = encode_image(image_bytes)
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a helpful assistant describe all text from given image."},
                {"role": "user", "content": [
                    {"type": "text", "text": "Describe the image all text and all objects in the image in detail and possible ?"},
                    {"type": "image_url", "image_url": {
                        "url": f"data:image/png;base64,{base64_image}"}
                    }
                ]}
            ],
            temperature=0.0,
        )
        result = response.choices[0].message.content
        # Create a Document object from the result text
        doc = LangchainDocument(page_content=result, metadata={"source": "image_analysis"})
        name_space = namespace_generator()
        try:
            pinecone_ar_class.insert_data_in_namespace(doc, embeddings, pinecone_index_name, name_space)
        except Exception as e:
            return {"output_result": f"An error occurred while processing the document: {str(e)}"}
        return {"output_result": name_space}
    except Exception as e:
        print(f"Error in image_qa: {str(e)}")
        return {"output_result": f"An error occurred while processing the image: {str(e)}"}

# file route type
def route_file_type(state):
    original_file_name = state["original_file_name"]
    file_extension = os.path.splitext(original_file_name)[1]
    if file_extension == ".png" or file_extension == ".jpg" or file_extension == ".jpeg":
        return "image_qa"
    elif file_extension == ".docx" or file_extension == ".doc":
        return "qa_retriever"
    elif file_extension == ".pdf":
        return "qa_retriever"
    elif file_extension == ".xlsx" or file_extension == ".csv":
        return "pandas_df_agent"
    else:
        return END


def graph_builder_function():    
    graph_builder = StateGraph(State)
    graph_builder.add_node("qa_retriever", QA_Retriever_node)
    graph_builder.add_node("pandas_df_agent", pandas_df_agent)
    graph_builder.add_node("image_qa", image_qa)
    graph_builder.set_conditional_entry_point(
    route_file_type,
    {
        "image_qa": "image_qa",
        "pandas_df_agent": "pandas_df_agent",
        "qa_retriever": "qa_retriever",
        }
    )
    # End each path appropriately
    graph_builder.add_edge("qa_retriever", END)
    graph_builder.add_edge("image_qa", END)
    graph_builder.add_edge("pandas_df_agent", END)
    # Compile the graph
    chatbot_graph = graph_builder.compile()
    return chatbot_graph


# graph_builder_fun = graph_builder_function()

# print(graph_builder_fun.get_graph().draw_mermaid_png())







def process_file_bytes(file_bytes: BytesIO, file_extension: str) -> str:
    """Process file bytes based on the file extension."""
    if file_extension == 'docx':
        return extract_text_from_docx(file_bytes)
    # Add more conditions for other file types if needed
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")

def extract_text_from_docx(file_bytes: BytesIO) -> str:
    """Extract text from a DOCX file."""
    doc = Document(file_bytes)
    docs = []
    for para in doc.paragraphs:
        tmp = para.text.replace('\n', ' ')
        tmp = re.sub(r'[‘”’“]', '', tmp)
        docs.append(tmp)
    
    return "\n".join(docs)