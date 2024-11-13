from langchain_core.documents import Document
# from langchain_community.document_loaders import Docx2txtLoader , PyMuPDFLoader , UnstructuredWordDocumentLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
import base64
import string , random , os
from datetime import datetime
import io
# from PIL import Image
from io import BytesIO
import re
pinecone_index_name = os.getenv("PINECONE_INDEX_NAME")

from docx import Document as DocxDocument
chunk_size = 800
class IncomingFileProcessor:
    def __init__(self, chunk_size=880) -> None:
        self.chunk_size = chunk_size
    def preprocess_files_all(self, raw_text:str, original_file_name: str):
        try:
            # Clean and process the raw text
            full_text = raw_text.replace('\n', ' ')
            full_text = re.sub(r'\s+', ' ', full_text)  # Replace multiple spaces with single space
            full_text = full_text.strip()
            
            if not full_text:
                raise ValueError("No text content found in the document")
            # Split the extracted text into chunks
            textsplit = RecursiveCharacterTextSplitter(
                chunk_size=self.chunk_size, chunk_overlap=15, length_function=len)
            doc_list = textsplit.split_documents([Document(page_content=full_text, metadata={"source": original_file_name})])
            return doc_list
        except Exception as e:
            print(f"Error processing DOCX file: {str(e)}")
            raise ValueError("An error occurred while processing the DOCX file.")
    

    # def get_pdf_splits(self, pdf_file: str,original_file_name: str):
    #     loader = PyMuPDFLoader(pdf_file)
    #     pages = loader.load()
    #     textsplit = RecursiveCharacterTextSplitter(
    #         chunk_size=self.chunk_size, chunk_overlap=15, length_function=len)
    #     doc_list = textsplit.split_documents(pages)
    #     # Create a dictionary to track page numbers and their content
    #     page_content_map = {}
    #     # Group chunks by page number
    #     for text in doc_list:
    #         page_num = text.metadata.get("page", 1)
    #         if page_num not in page_content_map:
    #             page_content_map[page_num] = []
    #         page_content_map[page_num].append(text.page_content)
        
    #     # Create new documents ensuring chunks from same page stay together
    #     final_docs = []
    #     for page_num, contents in page_content_map.items():
    #         combined_content = "\n".join(contents)
    #         final_docs.append(Document(
    #             page_content=combined_content,
    #             metadata={
    #                 "source": original_file_name,
    #                 "page": page_num
    #             }
    #         ))
            
    #     return final_docs
    
    # def get_doc_splits(self, doc_file: str,original_file_name: str):
    #     loader = Docx2txtLoader(str(doc_file))
    #     pages = loader.load()
    #     textsplit = RecursiveCharacterTextSplitter(
    #         chunk_size=self.chunk_size, chunk_overlap=15, length_function=len)
    #     doc_list = textsplit.split_documents(pages)
    #     # Create a dictionary to track page numbers and their content
    #     page_content_map = {}
    #     # Group chunks by page number
    #     for text in doc_list:
    #         page_num = text.metadata.get("page", 1)
    #         if page_num not in page_content_map:
    #             page_content_map[page_num] = []
    #         page_content_map[page_num].append(text.page_content)
        
    #     # Create new documents ensuring chunks from same page stay together
    #     final_docs = []
    #     for page_num, contents in page_content_map.items():
    #         combined_content = "\n".join(contents)
    #         final_docs.append(Document(
    #             page_content=combined_content,
    #             metadata={
    #                 "source": original_file_name,
    #                 "page": page_num
    #             }
    #         ))
            
    #     return final_docs  
    
    # def get_txt_splits(self, txt_file: str,original_file_name: str):
    #     with open(txt_file, 'r', encoding='utf-8') as txt_file:
    #         text_content = txt_file.read()
    #         texts = [Document(page_content=text_content, metadata={"source":original_file_name})]
    #     return texts


def encode_image(image_bytes: bytes) -> str:
    """Encode image bytes to a base64 string."""
    return base64.b64encode(image_bytes).decode("utf-8")

# def process_image(image_bytes: bytes):
#     """Process image bytes using PIL."""
#     with Image.open(io.BytesIO(image_bytes)) as img:
#         # Example: Convert image to grayscale
#         img = img.convert("L")
#         # Save or further process the image as needed
#         # For example, you can save it to a bytes buffer
#         buffer = io.BytesIO()
#         img.save(buffer, format="JPEG")
#         return buffer.getvalue()

def namespace_generator():
    curr_date = datetime.now().strftime("%Y-%m-%d-%H-%M")
    random_string = "".join(random.choices(string.ascii_letters, k=4))
    name_space = f"{random_string}{curr_date}:{pinecone_index_name}"
    return name_space






# @biochat_router.post("/retrieval")
# async def retrieval_file(query_input: QueryInput):
#     try:
#         query = query_input.query
#         collection_name = query_input.collection_name
#         if collection_name.endswith("xlsx") or collection_name.endswith("csv"):
#             if collection_name not in memory_storage:
#                 raise HTTPException(status_code=404, detail="File not found")
            
#             file_data = memory_storage[collection_name]
#             file_stream = file_data["content"]
#             original_filename = file_data["filename"]
#             file_stream.seek(0)
#             # Determine file type from filename
#             file_extension = original_filename.lower().split('.')[-1]
#             # Read file into DataFrame based on extension
#             if file_extension == 'csv':
#                 df = pd.read_csv(file_stream)
#             elif file_extension in ['xlsx', 'xls']:
#                 df = pd.read_excel(file_stream)   
#             pamdas_agent_result = pandas_df_agent(query,df)
#             return JSONResponse({"message": "Retrieval Successful!", "data": pamdas_agent_result}, status_code=status.HTTP_200_OK)
#     except Exception as ex:
#         return JSONResponse(content={"message": f"An error occurred: {str(ex)}"}, status_code=status.HTTP_500_INTERNAL_SERVER_ERROR)
